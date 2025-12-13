import os
import shutil
import tempfile
import logging

from pathlib import Path
from typing import List
from langchain_core.documents import Document
from .image_loader import load_image
from .logging_config import get_logger

from langchain_community.document_loaders import (
    PyPDFLoader,
    Docx2txtLoader,
    TextLoader,
)

logger = get_logger(__name__)


def _decode_text_bytes(data: bytes) -> str:
    # Try common encodings first (Windows txt often cp1252).
    for enc in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return data.decode(enc)
        except Exception:
            continue
    # Last resort: keep going with replacement chars.
    return data.decode("utf-8", errors="replace")


def _load_text_file(file_path: Path) -> List[Document]:
    data = file_path.read_bytes()
    text = _decode_text_bytes(data)
    return [Document(page_content=text, metadata={})]

def preprocess(docs: List[Document]) -> List[Document]:
    preprocessed_docs = []
    for doc in docs:
        text = doc.page_content
        text = text.strip()                # remove leading/trailing whitespace
        text = " ".join(text.split())      # normalize multiple spaces/newlines
        text = text.lower()                # lowercase for consistency        
        doc.page_content = text
        preprocessed_docs.append(doc)
    return preprocessed_docs

def load_documents(files: List, llm_model: str) -> List[Document]:
    logger.info("Loading %d files with LLM model %s", len(files), llm_model)

    # Support both Streamlit UploadedFile objects and filesystem paths.
    for f in files:
        if hasattr(f, "name"):
            logger.debug("File uploaded: %s", f.name)
        else:
            logger.debug("File path provided: %s", str(f))
    
    docs = []
    
    for file in files:
        if hasattr(file, "name"):
            original_filename = file.name  # Preserve original filename with spaces
            suffix = Path(original_filename).suffix.lower()
            raw_bytes = file.getbuffer()
        else:
            original_filename = str(file)
            suffix = Path(original_filename).suffix.lower()
            raw_bytes = None

        logger.debug("Processing file %s with suffix %s", original_filename, suffix)
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
            tmp_path = tmp_file.name
            if raw_bytes is not None:
                tmp_file.write(raw_bytes)

        if raw_bytes is None:
            # Copy local file content into temp file.
            try:
                shutil.copyfile(original_filename, tmp_path)
            except Exception:
                logger.exception("Error copying file to temp path: %s", original_filename)
                try:
                    os.unlink(tmp_path)
                except Exception:
                    pass
                continue

        try:
            file_path = Path(tmp_path)

            if suffix in [".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".webp", ".svg"]:
                documents = load_image(file_path, llm_model)
                if isinstance(documents, list):
                    docs.extend(documents)
                else:
                    docs.append(documents)
            else:            
                if suffix in [".docx", ".doc"]:
                    try:
                        # attempt to use the langchain loader (requires docx2txt)
                        loader = Docx2txtLoader(file_path)
                        documents = loader.load()
                    except Exception as e:
                        logger.debug("Docx2txtLoader failed: %s; falling back to python-docx", e)
                        try:
                            from docx import Document as _PyDocxDocument
                            _doc = _PyDocxDocument(str(file_path))
                            txt = "\n\n".join([p.text for p in _doc.paragraphs if p.text])
                            documents = [Document(page_content=txt, metadata={})]
                        except Exception as e2:
                            logger.exception("Failed to parse .docx file with both docx2txt and python-docx: %s", e2)
                            raise
                elif suffix in [".txt", ".md"]:
                    # TextLoader can fail on some Windows encodings; decode ourselves.
                    try:
                        documents = _load_text_file(file_path)
                    except Exception:
                        logger.exception("Failed to read text file: %s", original_filename)
                        raise
                elif suffix == ".pdf":
                    loader = PyPDFLoader(file_path)
                    documents = loader.load()

                # Update metadata for all documents in the list
                for doc in documents:
                    # Preserve original filename (may contain spaces)
                    doc.metadata = doc.metadata or {}
                    doc.metadata["source"] = original_filename
                    doc.metadata["path"] = original_filename
                    doc.metadata["file_ext"] = suffix
                    doc.metadata["file_type"] = "document"
                docs.extend(documents)
                
        finally:
            try:
                os.unlink(tmp_path)
            except Exception:
                pass

    if docs:
        docs = preprocess(docs)

    return docs