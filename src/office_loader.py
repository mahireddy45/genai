"""Office document loader (Word `.docx`)."""
from __future__ import annotations
from typing import List, Dict
from pathlib import Path
from docx import Document
from .schemas import DocumentMeta, IngestedDocument


def load_docx(path: str, chunk_size: int = 500, chunk_overlap: int = 50) -> List[Dict]:
    p = Path(path)
    if not p.is_file():
        return []

    doc = Document(path)
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text and p.text.strip())

    # simple chunking
    results = []
    start = 0
    length = len(full_text)
    chunk_index = 1
    while start < length:
        end = min(start + chunk_size, length)
        chunk = full_text[start:end].strip()
        if chunk:
            doc_id = f"{p.name}::c{chunk_index}"
            meta = DocumentMeta(filename=p.name, page=None, source=str(p), text_length=len(chunk))
            results.append(IngestedDocument(id=doc_id, text=chunk, meta=meta).dict())
            chunk_index += 1
        start = max(end - chunk_overlap, end)

    return results
